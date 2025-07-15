from docx import Document
import re
import pandas as pd
import os

def extrair_rdo(file_path):
    doc = Document(file_path)
    texto = "\n".join([p.text for p in doc.paragraphs])

    # Separar os relatórios diários pelo padrão "Relatório Diário de Operação"
    relatorios = re.split(r'\*Relatório Diário de Operação\*', texto)
    dados = []

    for rel in relatorios:
        if not rel.strip():
            continue

        # Extrair Data
        data_match = re.search(r'\*Data:\*\s*(\d{2}/\d{2}/\d{4})', rel)
        data = data_match.group(1) if data_match else None

        # Extrair Equipe (opcional)
        equipe_match = re.search(r'\*Equipe:\*\s*([^\n]+)', rel)
        equipe = equipe_match.group(1).strip() if equipe_match else None

        # Extrair Manutenção Corretiva (opcional)
        corretiva_match = re.search(r'\*Manutenção Corretiva\*\s*(.*?)\*Manutenção Preventiva\*', rel, re.DOTALL)
        corretiva = corretiva_match.group(1).strip() if corretiva_match else None

        # Extrair Manutenção Preventiva (opcional)
        preventiva_match = re.search(r'\*Manutenção Preventiva\*\s*(.*?)(\*Outras atividades\*|\*Status UFV\*|$)', rel, re.DOTALL)
        preventiva = preventiva_match.group(1).strip() if preventiva_match else None

        dados.append({
            'Data': data,
            'Equipe': equipe,
            'Manutenção Corretiva': corretiva,
            'Manutenção Preventiva': preventiva
        })

    df_rdo = pd.DataFrame(dados)
    return df_rdo

def extrair_contagem(file_path):
    doc = Document(file_path)
    texto = "\n".join([p.text for p in doc.paragraphs])

    # Extrair seção de falhas (entre "Falhas Julho 2025:" e "Ação:")
    falhas_match = re.search(r'Falhas Julho 2025:.*?(\d{2}):\s*(.*?)\nAção:', texto, re.DOTALL)
    if falhas_match:
        falhas_texto = texto.split("Falhas Julho 2025:")[1].split("Ação:")[0].strip()
    else:
        falhas_texto = ""

    # Outra estratégia: extrair linhas entre "Falhas Julho 2025:" e "Ação:"
    try:
        falhas_section = texto.split("Falhas Julho 2025:")[1].split("Ação:")[0]
    except IndexError:
        falhas_section = ""

    # Dividir por linhas e extrair itens tipo "Descrição - xQtd"
    falhas = []
    for line in falhas_section.splitlines():
        # Exemplo linha: "NCU perdeu comunicação (10 min) – x"
        line = line.strip()
        if not line:
            continue

        # Procurar padrão "descrição - xN" ou "descrição - N"
        m = re.match(r'(.*?)(?:[-–])\s*x?(\d+)', line)
        if m:
            descricao = m.group(1).strip()
            qtd = int(m.group(2))
            falhas.append({'Descrição': descricao, 'Quantidade': qtd})
        else:
            # Se não bateu, pode tentar extrair só a descrição sem quantidade
            falhas.append({'Descrição': line, 'Quantidade': None})

    df_contagem = pd.DataFrame(falhas)
    return df_contagem

if __name__ == "__main__":
    path_rdo = r"C:\Users\vinic\OneDrive - DSolar\Documentos\Performance\RDO\Julho RDO.docx"
    path_contagem = r"C:\Users\vinic\OneDrive - DSolar\Documentos\Performance\RDO\Contagem_ Julho.docx"

    df_rdo = extrair_rdo(path_rdo)
    df_contagem = extrair_contagem(path_contagem)

    print("=== Relatórios Diários de Operação (RDO) ===")
    print(df_rdo.head())

    print("\n=== Contagem de Falhas Julho ===")
    print(df_contagem.head())

    # Você pode salvar para CSV para análises futuras
    df_rdo.to_csv("rdo_extraido.csv", index=False)
    df_contagem.to_csv("contagem_falhas.csv", index=False)
